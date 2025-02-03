from docx import Document
import PyPDF2
import pdfplumber  # Advanced PDF text extraction
from pdf2image import convert_from_path  # PDF to image conversion
import pytesseract  # Optical Character Recognition (OCR)
import io
import re
import os
import json
import traceback
from typing import Optional, List, Dict, Any, Sequence, TypedDict, Annotated
import logging
from datetime import datetime
from dotenv import load_dotenv
import math
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder, PromptTemplate
from langchain_openai import ChatOpenAI
from langchain.vectorstores import FAISS
from langchain.embeddings import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.messages import AIMessage, BaseMessage, HumanMessage
from langgraph.checkpoint.memory import MemorySaver
from langgraph.graph import START, StateGraph
from langgraph.graph.message import add_messages
from django.apps import apps
from .models import ChatSession

from langchain.schema import BaseRetriever
from typing import List
from langchain_core.documents import Document
from .chainlogger import ChainLogger

from llama_parse import LlamaParse
from llama_index.core import SimpleDirectoryReader, VectorStoreIndex
from llama_index.core.schema import Document

from pydantic import Field
from langchain_core.retrievers import BaseRetriever
from langchain_core.documents import Document

from .processors.smart_document_processor import SmartDocumentProcessor
from pathlib import Path

import pymupdf4llm
import fitz  # PyMuPDF


logging.basicConfig(
    level=logging.DEBUG,  # Changed from INFO to DEBUG
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

try:
    pytesseract.get_tesseract_version()
    OCR_AVAILABLE = True
except Exception as e:
    OCR_AVAILABLE = False
    logger.warning(f"Tesseract OCR not available: {str(e)}")
    logger.warning("Some PDF processing features may be limited")

class ConversationState(TypedDict):
    input: str
    chat_history: Annotated[Sequence[BaseMessage], add_messages]
    context: str
    answer: str
    document_title: str 
    content_map: str
    metadata: dict
    
class ChatHistoryManager:
    def __init__(self):
        self._histories: dict[str, list[BaseMessage]] = {}
    
    def get_history(self, document_id: str) -> list[BaseMessage]:
        if document_id not in self._histories:
            self._histories[document_id] = []
        return self._histories[document_id]
    
    def add_message(self, document_id: str, message: BaseMessage):
        if document_id not in self._histories:
            self._histories[document_id] = []
        self._histories[document_id].append(message)
    
    def clear_history(self, document_id: str):
        if document_id in self._histories:
            self._histories[document_id] = []
    
    def get_all_histories(self) -> dict[str, list[BaseMessage]]:
        return self._histories

class LoggingVectorStoreRetriever(BaseRetriever):
    """Vector store retriever with enhanced logging."""
    
    vectorstore: Any = Field(None, description="Vector store instance")
    search_kwargs: Dict = Field(default_factory=lambda: {
        "k": 12,
        "fetch_k": 24,
        "score_threshold": 0.5
    })
    search_type: str = Field(default="similarity")
    verbose: bool = Field(default=True)

    class Config:
        """Pydantic config."""
        arbitrary_types_allowed = True

    def _get_relevant_documents(self, query: str) -> List[Document]:
        """Get relevant documents for a query with enhanced logging."""
        try:
            if self.search_type == "similarity":
                docs = self.vectorstore.similarity_search(
                    query, 
                    **self.search_kwargs
                )
            else:
                raise ValueError(f"Search type {self.search_type} not supported")
            
            # Log retrieval results
            if self.verbose:
                logger.info(f"\n{'='*80}\nQuery: {query}")
                logger.info(f"Retrieved {len(docs)} chunks with parameters: {self.search_kwargs}")
                for i, doc in enumerate(docs, 1):
                    logger.info(f"\nChunk {i}:")
                    logger.info(f"Content:\n{doc.page_content[:200]}...")  # Truncate long content
                    logger.info(f"Metadata: {doc.metadata}")
                    logger.info("-" * 40)
            
            return docs
            
        except Exception as e:
            logger.error(f"Error retrieving documents: {str(e)}")
            raise

def setup_retriever(vector_store: Any) -> LoggingVectorStoreRetriever:
    """Create a properly configured retriever."""
    search_kwargs = {
        "k": 16,
        "fetch_k": 24,
        "score_threshold": 0.5
    }
    
    try:
        retriever = LoggingVectorStoreRetriever(
            vectorstore=vector_store,
            search_kwargs=search_kwargs,
            verbose=True
        )
        
        # Verify configuration
        logger.info(f"Retriever configured with search_kwargs: {retriever.search_kwargs}")
        
        return retriever
        
    except Exception as e:
        logger.error(f"Error setting up retriever: {str(e)}")
        raise

# Create a global instance
chat_manager = ChatHistoryManager()

ChatSession = apps.get_model('app', 'ChatSession')

class LongDocumentProcessor:
    # Documents longer than 10,000 characters will be processed as long documents
    LONG_DOCUMENT_THRESHOLD = 400000  
    
    def __init__(self, max_chunk_size: int = 400000):  # 80000 chars per chunk
        self.max_chunk_size = max_chunk_size
        self.logger = logging.getLogger(__name__)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=16000,  # Smaller chunks for GPT processing
            chunk_overlap=200
        )
        self.llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.4)
    
    def split_document_by_size(self, text: str) -> List[str]:
        """Split document into chunks based on character size."""
        chunks = []
        current_chunk = []
        current_size = 0
        
        paragraphs = text.split('\n\n')
        
        for paragraph in paragraphs:
            paragraph_size = len(paragraph)
            
            if current_size + paragraph_size >= self.max_chunk_size:
                if current_chunk:
                    chunks.append('\n\n'.join(current_chunk))
                current_chunk = [paragraph]
                current_size = paragraph_size
            else:
                current_chunk.append(paragraph)
                current_size += paragraph_size
        
        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))
        
        return chunks

    def process_chunk(self, chunk: str, document_info: Dict[str, Any], chunk_index: int) -> str:
        """Process individual document chunk with detailed analysis."""
        try:
            prompt = f"""You are DiveInYT, a world-class teacher known for your ability to break down complex material 
into clear, engaging insights. One of your students has asked you to read this document 
and create detailed notes that will help him truly understand the material.

            As you read Segment {chunk_index + 1} of "{document_info['title']}", create the kind of notes 
you'd make for Student, knowing he wants to master this material.

            Please analyze this section as if you're explaining it to the student, covering:

1. Main Ideas Explained
- What's the big picture here, explained in a way that connects to what Dave already knows?
- What are the author's key arguments, and why do they matter?
- How does this fit into the larger context of the subject?

2. "Look at This!" Moments
- Choose 2-3 powerful quotes that really capture the essence of what's being taught
- For each quote, explain to Dave: "I chose this because..." and "This matters because..."
- Point out any "aha!" moments that Dave should pay special attention to

3. Building Your Vocabulary
- Which new terms or concepts would Dave need to understand?
- Explain each one using familiar examples or analogies
- Show how these terms connect to ideas Dave already knows

4. Making Connections
- How does this section build on what we've covered before?
- What foundations is it laying for what's coming next?
- What real-world applications or examples make this concrete?

5. Learning Check
- What are the most common misconceptions about these ideas?
- What questions should Dave be asking himself to check his understanding?
- What parts might need extra attention or review?

Text to Analyze:
{chunk}

Remember: You're not just summarizing - you're teaching how to think about and engage 
with these ideas. Use clear language, relevant examples, and helpful analogies. Point out 
patterns and connections that will help Dave build a deeper understanding.
"""
            
            response = self.llm.invoke(prompt)
            return response.content.strip()
            
        except Exception as e:
            self.logger.error(f"Error processing chunk {chunk_index}: {str(e)}")
            return f"Error processing chunk {chunk_index}"

    def merge_analyses(self, chunk_analyses: List[str], document_info: Dict[str, Any]) -> str:
        """Merge chunk analyses into a comprehensive document analysis."""
        try:
            combined_analyses = "\n\n".join([
                f"Segment {i+1}:\n{analysis}" 
                for i, analysis in enumerate(chunk_analyses)
            ])
            
            merge_prompt = f"""You are DiveInYT, and you've just finished reading "{document_info['title']}" by 
{document_info.get('author', 'the author')}. Now you're creating a comprehensive study guide 
for your student Dave, synthesizing all your insights into a clear, organized learning resource.

            Document Type: {document_info.get('file_type', 'Unknown')}

Create a thorough study guide that includes:

1. The Big Picture (Your Introduction to Dave)
"Dave, here's what this document is really about and why it matters..."
- Core message and significance
- How this connects to broader concepts you're learning
- Why understanding this will help you grow

2. Key Ideas and Themes
"Here are the main threads running through this work..."
- Major concepts and how they weave together
- Important patterns and relationships
- Questions these ideas help us answer

3. Learning Journey Map
"Let's see how the author builds their argument..."
- How the ideas develop and build on each other
- Critical turning points in the discussion
- How different parts work together

4. Essential Vocabulary and Concepts
"These are the tools you'll need in your mental toolkit..."
- Key terms explained clearly
- Core concepts unpacked
- Common confusions clarified

5. Deep Dive Topics
"Let's explore these particularly important areas..."
- Critical sections that need extra attention
- Complex ideas broken down step by step
- Connections between different concepts

6. Practical Applications
"Here's how these ideas work in the real world..."
- Real-world examples and applications
- How to use these concepts practically
- Connections to other things you're learning

7. Learning Checkpoints
"Test your understanding by..."
- Key questions to check comprehension
- Common misconceptions to watch for
- Areas that might need extra review

8. Next Steps
"To build on what you've learned..."
- Suggested follow-up readings
- Related topics to explore
- Skills to practice

Previous Section Notes:
{combined_analyses}

Remember: Your goal is to help Dave not just understand the content, but see how it fits into 
the bigger picture of what he's learning. Make connections explicit, anticipate questions, 
and provide clear examples that will help these ideas stick.
"""
            
            response = self.llm.invoke(merge_prompt)
            return response.content.strip()
            
        except Exception as e:
            self.logger.error(f"Error merging analyses: {str(e)}")
            return "Error creating full analysis"

    def process_long_document(self, text: str, document_info: Dict[str, Any]) -> str:
        """Main method to process long documents."""
        try:
            # Split into manageable chunks
            chunks = self.split_document_by_size(text)
            self.logger.info(f"Split document into {len(chunks)} chunks")
            
            # Process each chunk
            chunk_analyses = []
            for i, chunk in enumerate(chunks):
                self.logger.info(f"Processing chunk {i+1}/{len(chunks)}")
                chunk_analysis = self.process_chunk(chunk, document_info, i)
                chunk_analyses.append(chunk_analysis)
            
            # Merge analyses
            full_analysis = self.merge_analyses(chunk_analyses, document_info)
            
            return full_analysis
            
        except Exception as e:
            self.logger.error(f"Error processing long document: {str(e)}")
            raise

    def calculate_metrics(self, text: str) -> Dict[str, Any]:
        """Calculate document metrics like readability scores, complexity."""
        try:
            metrics = {
                'word_count': len(text.split()),
                'sentence_count': len(text.split('.')),
                'avg_sentence_length': len(text.split()) / max(len(text.split('.')), 1),
                # 'readability_score': calculate_readability(text),
                'technical_terms': self.extract_technical_terms(text)
            }
            return metrics
        except Exception as e:
            self.logger.error(f"Error calculating metrics: {str(e)}")
            return {}

    def extract_technical_terms(self, text: str) -> List[str]:
        """Extract and identify technical terminology from the text."""
        try:
            prompt = f"""Identify technical terms or specialized vocabulary in this text. 
            Return only the terms, one per line:

            {text[:2000]}  # Process first 2000 chars for term extraction
            """
            
            response = self.llm.invoke(prompt)
            terms = response.content.strip().split('\n')
            return [term.strip() for term in terms if term.strip()]
            
        except Exception as e:
            self.logger.error(f"Error extracting technical terms: {str(e)}")
            return []
        


def sanitize_filename(filename):
    return "".join([c for c in filename if c.isalpha() or c.isdigit() or c in ' -_.']).rstrip()

def create_embeddings():
    try:
        return OpenAIEmbeddings()
    except Exception as e:
        logger.error(f"Error creating embeddings: {str(e)}")
        raise

def load_vector_store(document_id):
    """Load vector store and create properly configured retriever."""
    vector_store_path = f"vector_stores/{document_id}"
    try:
        if os.path.exists(vector_store_path):
            embeddings = create_embeddings()
            vector_store = FAISS.load_local(
                vector_store_path, 
                embeddings, 
                allow_dangerous_deserialization=True
            )
            
            # Create retriever with explicit configuration
            retriever = setup_retriever(vector_store)
            
            return vector_store, retriever
            
        logger.warning(f"Vector store not found at path: {vector_store_path}")
        return None, None
        
    except Exception as e:
        logger.error(f"Error loading vector store for document ID {document_id}: {str(e)}")
        return None, None

def create_conversation_graph(retrieval_chain, document_title, content_map):
    def process_conversation(state: ConversationState):
        response = retrieval_chain.invoke({
            "input": state["input"],
            "chat_history": state["chat_history"],
            "document_title": state["document_title"],
            "content_map": state["content_map"]
        })
        
        return {
            "chat_history": [
                HumanMessage(content=state["input"]),
                AIMessage(content=response["answer"])
            ],
            "context": response.get("context", ""),
            "answer": response["answer"],
            "document_title": state["document_title"],
            "content_map": state["content_map"]
        }

    workflow = StateGraph(state_schema=ConversationState)
    workflow.add_node("conversation", process_conversation)
    workflow.add_edge(START, "conversation")
    
    memory = MemorySaver()
    return workflow.compile(checkpointer=memory)

document_summaries = {}

def store_summary(document_id, content_map):
    document_summaries[document_id] = {
        'content_map': content_map,
        'timestamp': datetime.now().isoformat()
    }
    logger.info(f"Stored content map for document ID {document_id}")
    
def store_summary_and_tags(document_id, summary_and_tags):
    if document_id not in document_summaries:
        document_summaries[document_id] = {}
    document_summaries[document_id]['summary_and_tags'] = summary_and_tags
    document_summaries[document_id]['updated_at'] = datetime.now().isoformat()
    logger.info(f"Stored summary and tags for document ID {document_id}")

def extract_text_from_word(file_path: str) -> Optional[str]:
    """Extract text content from Word documents."""
    try:
        doc = Document(file_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        for section in doc.sections:
            if section.header:
                header_text = section.header.paragraphs[0].text.strip()
                if header_text:
                    full_text.insert(0, f"Header: {header_text}")
            
            if section.footer:
                footer_text = section.footer.paragraphs[0].text.strip()
                if footer_text:
                    full_text.append(f"Footer: {footer_text}")
        
        processed_text = "\n\n".join(full_text)
        processed_text = re.sub(r'\s+', ' ', processed_text)
        processed_text = re.sub(r'\n\s*\n', '\n\n', processed_text)
        
        return processed_text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from Word document: {str(e)}")
        return None
        
        
        
import requests
import time
import os
import logging
import json
import base64
from typing import Optional, Dict, Any

logger = logging.getLogger(__name__)


class MuPDFDocumentProcessor:
    """
    A class to handle document processing using pymupdf4llm for converting documents to markdown.
    """
    def __init__(self, output_dir: str = "processed_documents"):
        """
        Initialize the MuPDF document processor.
        
        Args:
            output_dir (str): Directory to store processed documents and images
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
    def process_document(self, file_path: str, document_info: Dict[str, Any], pages: List[int] = None) -> Optional[str]:
        """
        Process a document using pymupdf4llm to convert it to markdown.
        Note: pymupdf4llm always processes all pages of the document.
        
        Args:
            file_path (str): Path to the document file
            document_info (dict): Information about the document
            pages (List[int], optional): This parameter is kept for interface compatibility
                                      but is not used by pymupdf4llm.
            
        Returns:
            Optional[str]: Processed markdown text or None if processing failed
        """
        try:
            if not self._validate_file(file_path):
                logger.error(f"Invalid or unsupported file: {file_path}")
                return None

            # Extract markdown content using pymupdf4llm
            # Note: the pages parameter is ignored as pymupdf4llm processes all pages by default
            if pages is not None:
                logger.warning("Page selection is not supported by pymupdf4llm. Processing all pages.")
            markdown_content = pymupdf4llm.to_markdown(file_path)
            
            # Extract metadata and images using PyMuPDF
            doc = fitz.open(file_path)
            metadata = self._extract_metadata(doc)
            
            # Extract and save images if needed
            images = self._extract_images(doc)
            if images:
                self._save_images(images, document_info['title'])
            
            # Add metadata header to markdown
            final_markdown = self._format_markdown(
                markdown_content,
                document_info,
                metadata,
                doc.page_count
            )
            
            doc.close()
            return final_markdown

        except Exception as e:
            logger.error(f"Error processing document with pymupdf4llm: {str(e)}")
            return None

    def _validate_file(self, file_path: str) -> bool:
        """Validate if the file exists and has a supported extension."""
        if not os.path.exists(file_path):
            return False
            
        supported_extensions = ['.pdf', '.epub', '.xps', '.oxps', '.cbz', '.fb2']
        return any(file_path.lower().endswith(ext) for ext in supported_extensions)

    def _extract_metadata(self, doc: fitz.Document) -> Dict:
        """Extract metadata from the document."""
        try:
            metadata = doc.metadata
            return {
                'author': metadata.get('author', 'Unknown'),
                'creation_date': self._parse_pdf_date(metadata.get('creationDate', '')),
                'modification_date': self._parse_pdf_date(metadata.get('modDate', '')),
                'producer': metadata.get('producer', 'Unknown'),
                'title': metadata.get('title', 'Unknown'),
                'subject': metadata.get('subject', 'Unknown'),
                'keywords': metadata.get('keywords', 'Unknown')
            }
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            return {}

    def _parse_pdf_date(self, date_str: str) -> str:
        """Parse PDF date format to readable string."""
        try:
            if not date_str:
                return "Unknown"
            
            # Remove D: prefix and timezone if present
            date_str = date_str.replace('D:', '')[:14]
            date_obj = datetime.strptime(date_str, '%Y%m%d%H%M%S')
            return date_obj.strftime('%Y-%m-%d %H:%M:%S')
        except Exception:
            return "Unknown"

    def _extract_images(self, doc: fitz.Document) -> Dict:
        """Extract images from the document."""
        images = {}
        try:
            for page_num, page in enumerate(doc):
                image_list = page.get_images()
                
                for img_idx, img in enumerate(image_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    
                    if base_image:
                        image_data = base_image["image"]
                        ext = base_image["ext"]
                        filename = f"page_{page_num + 1}_image_{img_idx + 1}.{ext}"
                        images[filename] = image_data
                        
        except Exception as e:
            logger.error(f"Error extracting images: {str(e)}")
        return images

    def _save_images(self, images: Dict, document_title: str) -> None:
        """Save extracted images from the document."""
        try:
            image_dir = os.path.join(self.output_dir, 'images', document_title)
            os.makedirs(image_dir, exist_ok=True)
            
            for filename, image_data in images.items():
                image_path = os.path.join(image_dir, filename)
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                    
            logger.info(f"Saved {len(images)} images to {image_dir}")
            
        except Exception as e:
            logger.error(f"Error saving images: {str(e)}")

    def _format_markdown(self, markdown: str, doc_info: Dict, metadata: Dict, page_count: int) -> str:
        """Format the markdown content with metadata."""
        metadata_header = f"""---
Title: {doc_info.get('title') or metadata.get('title', 'Unknown')}
Subject: {metadata.get('subject', 'Unknown')}
Keywords: {metadata.get('keywords', 'Unknown')}
Author: {metadata.get('author', 'Unknown')}
Created: {metadata.get('creation_date', 'Unknown')}
Modified: {metadata.get('modification_date', 'Unknown')}
Producer: {metadata.get('producer', 'Unknown')}
Pages: {page_count}
---

"""
        return metadata_header + markdown

# def batch_process_documents(input_dir: str, output_dir: str, pages: List[int] = None) -> Dict[str, Dict[str, Any]]:
#     """
#     Process multiple documents in a directory and generate a processing report.
    
#     Args:
#         input_dir (str): Directory containing input documents
#         output_dir (str): Directory to store processed documents
#         pages (List[int], optional): List of 0-based page numbers to process
        
#     Returns:
#         Dict[str, Dict[str, Any]]: Dictionary containing detailed processing results for each file
#     """
#     processor = MuPDFDocumentProcessor(output_dir)
#     results = {}
#     start_time = datetime.now()
    
#     # Create a directory for reports
#     report_dir = os.path.join(output_dir, 'reports')
#     os.makedirs(report_dir, exist_ok=True)
    
#     total_files = len([f for f in os.listdir(input_dir) 
#                       if processor._validate_file(os.path.join(input_dir, f))])
#     processed_files = 0
#     failed_files = 0
    
#     for file_name in os.listdir(input_dir):
#         file_path = os.path.join(input_dir, file_name)
#         if processor._validate_file(file_path):
#             file_start_time = datetime.now()
            
#             try:
#                 document_info = {'title': os.path.splitext(file_name)[0]}
#                 markdown_content = processor.process_document(file_path, document_info, pages)
                
#                 if markdown_content:
#                     output_path = os.path.join(output_dir, f"{document_info['title']}.md")
#                     Path(output_path).write_text(markdown_content, encoding='utf-8')
#                     processed_files += 1
                    
#                     # Store detailed success information
#                     results[file_name] = {
#                         'status': 'success',
#                         'processing_time': str(datetime.now() - file_start_time),
#                         'output_path': output_path,
#                         'pages_processed': len(pages) if pages else 'all',
#                         'output_size': len(markdown_content),
#                         'timestamp': datetime.now().isoformat()
#                     }
#                 else:
#                     failed_files += 1
#                     results[file_name] = {
#                         'status': 'failed',
#                         'error': 'No content generated',
#                         'timestamp': datetime.now().isoformat()
#                     }
                    
#             except Exception as e:
#                 failed_files += 1
#                 error_msg = str(e)
#                 logger.error(f"Error processing {file_name}: {error_msg}")
#                 results[file_name] = {
#                     'status': 'failed',
#                     'error': error_msg,
#                     'timestamp': datetime.now().isoformat()
#                 }
    
#     # Generate and save processing report
#     total_time = datetime.now() - start_time
#     report = {
#         'summary': {
#             'total_files': total_files,
#             'processed_files': processed_files,
#             'failed_files': failed_files,
#             'success_rate': f"{(processed_files/total_files)*100:.2f}%" if total_files > 0 else "0%",
#             'total_processing_time': str(total_time),
#             'average_time_per_file': str(total_time / total_files) if total_files > 0 else "N/A",
#             'timestamp': datetime.now().isoformat()
#         },
#         'file_results': results
#     }
    
#     # Save report as JSON
#     report_path = os.path.join(report_dir, f'processing_report_{start_time.strftime("%Y%m%d_%H%M%S")}.json')
#     Path(report_path).write_text(json.dumps(report, indent=2), encoding='utf-8')
    
#     # Generate markdown report
#     md_report = f"""# Document Processing Report

# ## Summary
# - Total Files: {report['summary']['total_files']}
# - Successfully Processed: {report['summary']['processed_files']}
# - Failed: {report['summary']['failed_files']}
# - Success Rate: {report['summary']['success_rate']}
# - Total Processing Time: {report['summary']['total_processing_time']}
# - Average Time per File: {report['summary']['average_time_per_file']}

# ## Detailed Results

# | File | Status | Processing Time | Output Size | Error |
# |------|--------|----------------|-------------|-------|
# """
    
#     for file_name, result in results.items():
#         status = result['status']
#         time = result.get('processing_time', 'N/A')
#         size = result.get('output_size', 'N/A')
#         error = result.get('error', 'None')
#         md_report += f"| {file_name} | {status} | {time} | {size} | {error} |\n"

#     # Save markdown report
#     md_report_path = os.path.join(report_dir, f'processing_report_{start_time.strftime("%Y%m%d_%H%M%S")}.md')
#     Path(md_report_path).write_text(md_report, encoding='utf-8')
    
#     return results

class PDFDocumentProcessor:
    """
    A class to handle PDF document processing using pdfplumber for converting PDFs to markdown.
    """
    def __init__(self, output_dir: str = "processed_documents"):
        """
        Initialize the PDF document processor.
        
        Args:
            output_dir (str): Directory to store processed documents and images
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
    def save_processed_markdown(document_id: str, markdown_text: str, document_info: dict = None) -> str:
        """
        Save processed markdown text to a directory for inspection.
        
        Args:
            document_id: Unique identifier for the document
            markdown_text: The processed markdown content
            document_info: Optional metadata about the document
            
        Returns:
            str: Path to the saved markdown file
        """
        try:
            # Create base directory for processed documents
            base_dir = "processed_documents"
            os.makedirs(base_dir, exist_ok=True)
            
            # Create timestamped directory for this processing run
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            doc_dir = os.path.join(base_dir, f"{document_id}_{timestamp}")
            os.makedirs(doc_dir, exist_ok=True)
            
            # Save markdown content
            markdown_path = os.path.join(doc_dir, "processed_content.md")
            with open(markdown_path, 'w', encoding='utf-8') as f:
                # Add metadata header if available
                if document_info:
                    f.write("---\n")
                    for key, value in document_info.items():
                        f.write(f"{key}: {value}\n")
                    f.write("---\n\n")
                
                # Write main content
                f.write(markdown_text)
            
            # Save processing info
            info_path = os.path.join(doc_dir, "processing_info.txt")
            with open(info_path, 'w', encoding='utf-8') as f:
                f.write(f"Processing timestamp: {timestamp}\n")
                f.write(f"Document ID: {document_id}\n")
                if document_info:
                    f.write("\nDocument Information:\n")
                    for key, value in document_info.items():
                        f.write(f"{key}: {value}\n")
            
            logger.info(f"Saved processed markdown to: {markdown_path}")
            logger.info(f"Saved processing info to: {info_path}")
            
            return markdown_path
            
        except Exception as e:
            logger.error(f"Error saving processed markdown: {str(e)}")
            raise
        
    def process_document(self, file_path: str, document_info: Dict[str, Any]) -> Optional[str]:
        """
        Process a PDF document using pdfplumber to convert it to markdown.
        
        Args:
            file_path (str): Path to the PDF file
            document_info (dict): Information about the document
            
        Returns:
            Optional[str]: Processed markdown text or None if processing failed
        """
        try:
            if not self._validate_pdf(file_path):
                logger.error(f"Invalid or unsupported file: {file_path}")
                return None

            markdown_content = []
            images = {}
            metadata = {}
            
            with pdfplumber.open(file_path) as pdf:
                # Extract PDF metadata
                metadata = self._extract_metadata(pdf)
                
                # Process each page
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text()
                    if text:
                        markdown_content.append(f"\n## Page {page_num}\n\n{text}")
                    
                    # Extract images
                    page_images = self._extract_images(page, page_num)
                    images.update(page_images)
                
                # Save images if any were extracted
                if images:
                    self._save_images(images, document_info['title'])
                
                # Combine all content and add metadata
                final_markdown = self._format_markdown(
                    "\n".join(markdown_content),
                    document_info,
                    metadata,
                    len(pdf.pages)
                )
                
                return final_markdown

        except Exception as e:
            logger.error(f"Error processing PDF with pdfplumber: {str(e)}")
            return None

    def _validate_pdf(self, file_path: str) -> bool:
        """Validate if the file is a PDF."""
        if not os.path.exists(file_path):
            return False
        return file_path.lower().endswith('.pdf')

    def _extract_metadata(self, pdf) -> Dict:
        """Extract metadata from the PDF."""
        try:
            doc_info = pdf.doc.info[0]
            return {
                'author': doc_info.get('/Author', b'Unknown').decode('utf-8', errors='ignore'),
                'creation_date': self._parse_pdf_date(doc_info.get('/CreationDate', b'')),
                'modification_date': self._parse_pdf_date(doc_info.get('/ModDate', b'')),
                'producer': doc_info.get('/Producer', b'Unknown').decode('utf-8', errors='ignore')
            }
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            return {}

    def _parse_pdf_date(self, date_str: bytes) -> str:
        """Parse PDF date format to readable string."""
        try:
            if not date_str:
                return "Unknown"
            
            date_str = date_str.decode('utf-8', errors='ignore')
            # Remove D: prefix and timezone if present
            date_str = date_str.replace('D:', '')[:14]
            date_obj = datetime.strptime(date_str, '%Y%m%d%H%M%S')
            return date_obj.strftime('%Y-%m-%d %H:%M:%S')
        except Exception:
            return "Unknown"

    def _extract_images(self, page, page_num: int) -> Dict:
        """Extract images from a PDF page."""
        images = {}
        try:
            # Extract images using pdfplumber's image extraction
            for image_num, image in enumerate(page.images, 1):
                if 'stream' in image:
                    filename = f"page_{page_num}_image_{image_num}.png"
                    images[filename] = image['stream'].get_data()
        except Exception as e:
            logger.error(f"Error extracting images from page {page_num}: {str(e)}")
        return images

    def _save_images(self, images: Dict, document_title: str) -> None:
        """Save extracted images from the document."""
        try:
            image_dir = os.path.join(self.output_dir, 'images', document_title)
            os.makedirs(image_dir, exist_ok=True)
            
            for filename, image_data in images.items():
                image_path = os.path.join(image_dir, filename)
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                    
            logger.info(f"Saved {len(images)} images to {image_dir}")
            
        except Exception as e:
            logger.error(f"Error saving images: {str(e)}")

    def _format_markdown(self, markdown: str, doc_info: Dict, metadata: Dict, page_count: int) -> str:
        """Format the markdown content with metadata."""
        metadata_header = f"""---
Title: {doc_info['title']}
File Type: PDF
Pages: {page_count}
Created: {metadata.get('creation_date', 'Unknown')}
Modified: {metadata.get('modification_date', 'Unknown')}
Author: {metadata.get('author', 'Unknown')}
Producer: {metadata.get('producer', 'Unknown')}
---

"""
        return metadata_header + markdown

class MarkerDocumentProcessor:
    """
    A class to handle document processing using the Marker API for converting documents to markdown.
    """
    def __init__(self, api_key: str, base_url: str = "https://www.datalab.to/api/v1"):
        """
        Initialize the Marker document processor.
        
        Args:
            api_key (str): The API key for authentication
            base_url (str): Base URL for the Marker API
        """
        self.api_key = api_key
        self.base_url = base_url
        self.headers = {"X-Api-Key": api_key}
        
    def process_document(self, file_path: str, document_info: Dict[str, Any]) -> Optional[str]:
        """
        Process a document using Marker API to convert it to markdown.
        
        Args:
            file_path (str): Path to the document file
            document_info (dict): Information about the document
            
        Returns:
            Optional[str]: Processed markdown text or None if processing failed
        """
        try:
            # Determine file type and mime type
            mime_type = self._get_mime_type(file_path)
            if not mime_type:
                logger.error(f"Unsupported file type for file: {file_path}")
                return None

            # Prepare the form data for the API request
            form_data = {
                'file': (os.path.basename(file_path), open(file_path, 'rb'), mime_type),
                'langs': (None, "English"),  # Default to English, could be made configurable
                'output_format': (None, 'markdown'),
                'force_ocr': (None, False),
                'paginate': (None, True),  # Enable pagination for better structure
                'use_llm': (None, True),   # Enable LLM for better accuracy
                'strip_existing_ocr': (None, False),
            }

            # Make the initial API request
            response = self._make_api_request("marker", form_data)
            if not response or not response.get('success'):
                logger.error(f"Initial API request failed: {response.get('error', 'Unknown error')}")
                return None

            # Poll for results
            result = self._poll_for_results(response['request_check_url'])
            if not result or not result.get('success'):
                logger.error(f"Failed to get results: {result.get('error', 'Unknown error')}")
                return None

            # Process and save images if any
            if result.get('images'):
                self._save_images(result['images'], document_info['title'])

            # Add metadata to markdown
            markdown_content = self._format_markdown(
                result['markdown'], 
                document_info,
                result.get('metadata', {}),
                result.get('page_count', 0)
            )

            return markdown_content

        except Exception as e:
            logger.error(f"Error processing document with Marker: {str(e)}")
            return None

    def _make_api_request(self, endpoint: str, form_data: Dict) -> Optional[Dict]:
        """Make an API request to the Marker service."""
        try:
            url = f"{self.base_url}/{endpoint}"
            response = requests.post(url, files=form_data, headers=self.headers)
            response.raise_for_status()
            return response.json()
        except Exception as e:
            logger.error(f"API request failed: {str(e)}")
            return None

    def _poll_for_results(self, check_url: str, max_polls: int = 300, poll_interval: int = 2) -> Optional[Dict]:
        """Poll for results from the Marker API."""
        for _ in range(max_polls):
            try:
                time.sleep(poll_interval)
                response = requests.get(check_url, headers=self.headers)
                response.raise_for_status()
                data = response.json()
                
                if data["status"] == "complete":
                    return data
                    
            except Exception as e:
                logger.error(f"Error polling for results: {str(e)}")
                return None
                
        logger.error("Maximum polling attempts reached")
        return None

    def _save_images(self, images: Dict, document_title: str) -> None:
        """Save extracted images from the document."""
        try:
            image_dir = f"processed_documents/images/{document_title}"
            os.makedirs(image_dir, exist_ok=True)
            
            for filename, image_data in images.items():
                image_path = os.path.join(image_dir, filename)
                image_bytes = base64.b64decode(image_data)
                
                with open(image_path, 'wb') as f:
                    f.write(image_bytes)
                    
            logger.info(f"Saved {len(images)} images to {image_dir}")
            
        except Exception as e:
            logger.error(f"Error saving images: {str(e)}")

    def _format_markdown(self, markdown: str, doc_info: Dict, metadata: Dict, page_count: int) -> str:
        """Format the markdown content with metadata."""
        metadata_header = f"""---
Title: {doc_info['title']}
File Type: {doc_info.get('file_type', 'Unknown')}
Pages: {page_count}
Created: {metadata.get('creation_date', 'Unknown')}
Modified: {metadata.get('modification_date', 'Unknown')}
Author: {metadata.get('author', 'Unknown')}
---

"""
        return metadata_header + markdown

    def _get_mime_type(self, file_path: str) -> Optional[str]:
        """Get the mime type for a file."""
        try:
            import filetype
            guess = filetype.guess(file_path)
            if guess:
                return guess.mime
            
            # Fallback to extension-based detection
            ext = os.path.splitext(file_path)[1].lower()
            mime_types = {
                '.pdf': 'application/pdf',
                '.doc': 'application/msword',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.ppt': 'application/vnd.ms-powerpoint',
                '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
            }
            return mime_types.get(ext)
            
        except ImportError:
            logger.warning("filetype package not installed, using basic extension check")
            ext = os.path.splitext(file_path)[1].lower()
            return self._get_basic_mime_type(ext)
        except Exception as e:
            logger.error(f"Error determining mime type: {str(e)}")
            return None

    def _get_basic_mime_type(self, extension: str) -> Optional[str]:
        """Basic mime type determination based on file extension."""
        mime_map = {
            '.pdf': 'application/pdf',
            '.doc': 'application/msword',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.ppt': 'application/vnd.ms-powerpoint',
            '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
        }
        return mime_map.get(extension.lower())

def get_document_info(file_path: str) -> dict:
    """
    Get document info with enhanced document name extraction and metadata handling.
    
    Args:
        file_path (str): Path to the document
        
    Returns:
        dict: Dictionary containing document information with properly extracted title and metadata
    """
    try:
        # Validate file existence
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")

        file_name = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        modification_time = os.path.getmtime(file_path)
        
        # Initialize document title and metadata
        document_title = None
        metadata = {}

        # Try to get title and metadata from PDF metadata first
        if file_path.lower().endswith('.pdf'):
            # Try pdfplumber first
            try:
                with pdfplumber.open(file_path) as pdf:
                    if pdf.metadata:
                        document_title = pdf.metadata.get('Title', '').strip()
                        metadata.update({
                            'author': pdf.metadata.get('Author', 'Unknown'),
                            'creation_date': pdf.metadata.get('CreationDate', 'Unknown'),
                            'modification_date': pdf.metadata.get('ModDate', 'Unknown'),
                            'subject': pdf.metadata.get('Subject', 'Unknown'),
                            'producer': pdf.metadata.get('Producer', 'Unknown')
                        })
            except Exception as e:
                logger.warning(f"Failed to extract PDF metadata using pdfplumber: {str(e)}")
            finally:
                # Fallback to PyPDF2 if pdfplumber doesn't get a title
                if not document_title:
                    try:
                        with open(file_path, 'rb') as file:
                            pdf_reader = PyPDF2.PdfReader(file, strict=False)
                            if pdf_reader.metadata:
                                document_title = pdf_reader.metadata.get('/Title', '').strip()
                                metadata.update({
                                    'author': pdf_reader.metadata.get('/Author', 'Unknown'),
                                    'creation_date': pdf_reader.metadata.get('/CreationDate', 'Unknown'),
                                    'modification_date': pdf_reader.metadata.get('/ModDate', 'Unknown'),
                                    'subject': pdf_reader.metadata.get('/Subject', 'Unknown'),
                                    'producer': pdf_reader.metadata.get('/Producer', 'Unknown')
                                })
                    except Exception as e:
                        logger.warning(f"Failed to extract PDF metadata with PyPDF2: {str(e)}")
                    finally:
                        if not document_title:
                            logger.info("No title found in PDF metadata, will use filename")

        # For Word documents, try to get title from core properties
        elif file_path.lower().endswith(('.doc', '.docx')):
            try:
                doc = Document(file_path)
                if doc.core_properties.title:
                    document_title = doc.core_properties.title.strip()
                metadata.update({
                    'author': doc.core_properties.author or 'Unknown',
                    'creation_date': doc.core_properties.created.isoformat() if doc.core_properties.created else 'Unknown',
                    'modification_date': doc.core_properties.modified.isoformat() if doc.core_properties.modified else 'Unknown',
                    'subject': doc.core_properties.subject or 'Unknown'
                })
            except Exception as e:
                logger.warning(f"Failed to extract Word document metadata: {str(e)}")
            finally:
                if not document_title:
                    logger.info("No title found in Word document properties, will use filename")

        # If no title found in metadata, use filename but clean it up
        try:
            if not document_title:
                # Remove file extension and replace underscores/hyphens with spaces
                document_title = os.path.splitext(file_name)[0]
                document_title = document_title.replace('_', ' ').replace('-', ' ')
                # Remove common prefixes/suffixes like 'copy', 'final', 'v1', etc.
                document_title = re.sub(r'(?i)\b(copy|final|draft|v\d+|rev\d+|version\d+|doc\d+|report\d+)\b', '', document_title)
                # Remove dates and timestamps (e.g., 2023-10-01, 20231001)
                document_title = re.sub(r'\b\d{4}[-_]?\d{2}[-_]?\d{2}\b', '', document_title)
                # Clean up extra spaces and capitalize properly
                document_title = ' '.join(document_title.split()).strip()
                # Capitalize first letter of each word
                document_title = document_title.title()
        except Exception as e:
            logger.error(f"Error cleaning up filename for title: {str(e)}")
            document_title = os.path.splitext(file_name)[0]  # Fallback to basic filename

        return {
            'title': document_title,
            'file_type': os.path.splitext(file_name)[1][1:].lower(),
            'file_size': file_size,
            'modification_date': datetime.fromtimestamp(modification_time).isoformat(),
            'metadata': metadata
        }
        
    except Exception as e:
        logger.error(f"Error getting document information: {str(e)}")
        return {
            'title': os.path.splitext(os.path.basename(file_path))[0],
            'file_type': os.path.splitext(file_path)[1][1:].lower() if '.' in file_path else '',
            'file_size': 0,
            'modification_date': datetime.now().isoformat(),
            'metadata': {}
        }

def check_pdf_dependencies() -> dict:
    """
    Check if all required PDF processing dependencies are available
    
    Returns:
        dict: Status of each dependency
    """
    status = {
        'poppler': False,
        'tesseract': False,
        'pdfplumber': False
    }
    
    # Check Poppler (required for pdf2image)
    try:
        from pdf2image.pdf2image import get_poppler_path
        poppler_path = get_poppler_path()
        if poppler_path:
            status['poppler'] = True
    except Exception as e:
        logger.warning(f"Poppler not properly configured: {str(e)}")
    
    # Check Tesseract
    try:
        pytesseract.get_tesseract_version()
        status['tesseract'] = True
    except Exception as e:
        logger.warning(f"Tesseract not properly configured: {str(e)}")
    
    # Check pdfplumber
    try:
        import pdfplumber
        status['pdfplumber'] = True
    except ImportError:
        logger.warning("pdfplumber not installed")
    
    return status

def get_document_metadata(file_path: str) -> dict:
    """Extract detailed metadata from documents."""
    try:
        metadata = {}
        
        if file_path.endswith('.pdf'):
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata = {
                    'num_pages': len(pdf_reader.pages),
                    'pdf_info': pdf_reader.metadata or {}
                }
                
        elif file_path.endswith(('.doc', '.docx')):
            doc = Document(file_path)
            metadata = {
                'num_pages': len(doc.sections),
                'num_paragraphs': len(doc.paragraphs),
                'core_properties': {
                    'author': doc.core_properties.author,
                    'created': doc.core_properties.created,
                    'modified': doc.core_properties.modified,
                    'title': doc.core_properties.title
                }
            }
            
        return metadata
        
    except Exception as e:
        logger.error(f"Error extracting document metadata: {str(e)}")
        return {}


# Global cache to track document processing status
document_processing_cache = {}

def is_document_processed(document_id: str) -> bool:
    """Check if a document has already been processed."""
    return document_id in document_processing_cache

def mark_document_processed(document_id: str):
    """Mark a document as processed."""
    document_processing_cache[document_id] = {
        'processed': True,
        'timestamp': datetime.now().isoformat()
    }

def save_processed_markdown(document_id: str, markdown_text: str, document_info: dict = None) -> str:
    """
    Save processed markdown text to a directory for inspection.
    
    Args:
        document_id: Unique identifier for the document
        markdown_text: The processed markdown content
        document_info: Optional metadata about the document
        
    Returns:
        str: Path to the saved markdown file
    """
    try:
        # Create base directory for processed documents
        base_dir = "processed_documents"
        os.makedirs(base_dir, exist_ok=True)
        
        # Create timestamped directory for this processing run
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_dir = os.path.join(base_dir, f"{document_id}_{timestamp}")
        os.makedirs(doc_dir, exist_ok=True)
        
        # Save markdown content
        markdown_path = os.path.join(doc_dir, "processed_content.md")
        with open(markdown_path, 'w', encoding='utf-8') as f:
            # Add metadata header if available
            if document_info:
                f.write("---\n")
                for key, value in document_info.items():
                    f.write(f"{key}: {value}\n")
                f.write("---\n\n")
            
            # Write main content
            f.write(markdown_text)
        
        # Save processing info
        info_path = os.path.join(doc_dir, "processing_info.txt")
        with open(info_path, 'w', encoding='utf-8') as f:
            f.write(f"Processing timestamp: {timestamp}\n")
            f.write(f"Document ID: {document_id}\n")
            if document_info:
                f.write("\nDocument Information:\n")
                for key, value in document_info.items():
                    f.write(f"{key}: {value}\n")
        
        logger.info(f"Saved processed markdown to: {markdown_path}")
        logger.info(f"Saved processing info to: {info_path}")
        
        return markdown_path
        
    except Exception as e:
        logger.error(f"Error saving processed markdown: {str(e)}")
        raise


# def process_document(document_id, file_path):
#     """Process a document with enhanced LLM-based TOC extraction and optimized summary generation."""
#     logger.info(f"Starting to process document ID: {document_id}")
    
#     try:
#         # Get document info
#         document_info = get_document_info(file_path)
#         if not document_info:
#             logger.error(f"Failed to get document info for document ID: {document_id}")
#             return None, None, None, None, None, None, None

#         # Extract text from document using MuPDF
#         processor = MuPDFDocumentProcessor(output_dir="processed_documents")
#         markdown_content = processor.process_document(file_path, document_info)
#         if not markdown_content:
#             logger.error(f"Failed to extract text from document ID: {document_id}")
#             return None, None, None, None, None, None, None

#         logger.info(f"Text extracted from document ID: {document_id}. Length: {len(markdown_content)}")
        
#         try:
#             saved_path = save_processed_markdown(
#                 document_id=document_id,
#                 markdown_text=markdown_content,
#                 document_info=document_info
#             )
#             logger.info(f"Saved processed markdown to: {saved_path}")
#         except Exception as e:
#             logger.error(f"Error saving markdown: {str(e)}")

#         try:
#             # Initialize SmartDocumentProcessor with LLM client
#             llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.3)
#             doc_processor = SmartDocumentProcessor()
            
#             # Process document and extract TOC using LLM
#             processed_chunks, content_map = doc_processor.process_document(markdown_content)
            
#             # Add validation for processed_chunks
#             if not processed_chunks:
#                 logger.warning("No chunks generated, creating fallback chunk")
#                 processed_chunks = [{
#                     'content': "Document content unavailable. Please check if:\n"
#                             "1. The document has readable text\n"
#                             "2. Contains more than 50 words\n"
#                             "3. Isn't image-based",
#                     'metadata': {'fallback': True}
#                 }]
            
#             # Debug logging for chunks
#             logger.debug(f"Number of chunks created: {len(processed_chunks)}")
#             if processed_chunks:
#                 logger.debug(f"Sample chunk structure: {processed_chunks[0].keys()}")
            
#             # Create and save vector store
#             try:
#                 embeddings = create_embeddings()
                
#                 # Prepare chunks for vector store
#                 texts = []
#                 metadatas = []
                
#                 for i, chunk in enumerate(processed_chunks):
#                     if not isinstance(chunk, dict) or 'content' not in chunk:
#                         logger.error(f"Invalid chunk format at index {i}: {chunk}")
#                         continue
                    
#                     texts.append(chunk['content'])
#                     metadata = {
#                         "document_id": document_id,
#                         "chunk_index": i,
#                         "total_chunks": len(processed_chunks)
#                     }
                    
#                     # Add section metadata from LLM-extracted TOC
#                     if 'metadata' in chunk and isinstance(chunk['metadata'], dict):
#                         for key, value in chunk['metadata'].items():
#                             if value is not None:
#                                 metadata[key] = value
                    
#                     metadatas.append(metadata)
                
#                 logger.info(f"Creating vector store with {len(texts)} chunks")
                
#                 if not processed_chunks:
#                     logger.error("No valid chunks created. Possible causes:")
#                     logger.error("- Document structure not recognized")
#                     logger.error("- Overly strict content filters")
#                     logger.error("- Empty/malformed document")
                    
#                     # Create fallback chunk
#                     chunks = [{
#                         'content': "Document content unavailable. Please check if:\n"
#                                 "1. The document has readable text\n"
#                                 "2. Contains more than 50 words\n"
#                                 "3. Isn't image-based",
#                         'metadata': {'fallback': True}
#                     }]
                
#                 # Create vector store with enhanced metadata
#                 vector_store = FAISS.from_texts(
#                     texts,
#                     embeddings,
#                     metadatas=metadatas
#                 )
                
#                 # Create retriever with optimized parameters
#                 retriever = LoggingVectorStoreRetriever(
#                     vectorstore=vector_store,
#                     search_kwargs={
#                         "k": 4,
#                         "fetch_k": 24,
#                         "score_threshold": 0.5
#                     }
#                 )
                
#                 # Save vector store
#                 vector_store_path = f"vector_stores/{document_id}"
#                 os.makedirs(os.path.dirname(vector_store_path), exist_ok=True)
#                 vector_store.save_local(vector_store_path)
#                 logger.info(f"Created and saved vector store for document ID: {document_id}")
                
#                 # Set up retrieval chain
#                 retrieval_chain = setup_llm_chain(retriever)
                
#                 # Store content map
#                 store_summary(document_id, content_map)
                
#                 # Generate optimized summary using only content map
#                 try:
#                     # Create a focused summary prompt
#                     summary_prompt = f"""Analyze the structure and organization of this document based on its content map:

# Document Title: {document_info['title']}

# Content Map Structure:
# {json.dumps(content_map, indent=2)}

# Please provide:
# 1. Key points and main ideas evident from the document structure
# 2. Major themes and concepts based on the section organization
# 3. Document's organizational structure and hierarchy
# 4. Relevant keywords or tags derived from section titles

# Focus your analysis on the organizational structure shown in the content map."""

#                     # Create minimal initial state
#                     initial_state = {
#                         "input": summary_prompt,
#                         "chat_history": [],
#                         "document_title": document_info['title'],
#                         "content_map": content_map
#                     }
                    
#                     # Create conversation graph with minimal configuration
#                     conversation_app = create_conversation_graph(
#                         retrieval_chain,
#                         document_info['title'],
#                         content_map
#                     )
                    
#                     config = {
#                         "configurable": {
#                             "thread_id": f"document_{document_id}",
#                             "checkpoint_id": f"initial_summary_{document_id}",
#                             "checkpoint_ns": "document_processing"
#                         }
#                     }
                    
#                     summary_result = conversation_app.invoke(
#                         initial_state,
#                         config=config
#                     )
                    
#                     summary_and_tags = summary_result["answer"]
#                     store_summary_and_tags(document_id, summary_and_tags)
#                     mark_document_processed(document_id)
                    
#                     return (vector_store, retrieval_chain, document_info, 
#                             content_map, summary_and_tags, [], markdown_content)
            
#                 except Exception as e:
#                     logger.error(f"Error generating summary: {str(e)}")
#                     return (vector_store, retrieval_chain, document_info, 
#                             content_map, "", [], markdown_content)
                        
#             except Exception as e:
#                 logger.error(f"Error creating vector store: {str(e)}")
#                 logger.error(f"Chunk format: {processed_chunks[0] if processed_chunks else 'No chunks'}")
#                 raise
                
#         except Exception as e:
#             logger.error(f"Error in document processing: {str(e)}")
#             return None, None, None, None, None, None, None

#     except Exception as e:
#         logger.error(f"Unexpected error processing document ID {document_id}: {str(e)}")
#         logger.error(traceback.format_exc())
#         return None, None, None, None, None, None, None

def process_document(document_id, file_path):
    """Process a document with fixed chunk creation and vector store alignment."""
    logger.info(f"Starting to process document ID: {document_id}")
    
    try:
        # Get document info
        document_info = get_document_info(file_path)
        if not document_info:
            logger.error(f"Failed to get document info for document ID: {document_id}")
            return None, None, None, None, None, None, None

        # Extract text from document using MuPDF
        processor = MuPDFDocumentProcessor(output_dir="processed_documents")
        markdown_content = processor.process_document(file_path, document_info)
        if not markdown_content:
            logger.error(f"Failed to extract text from document ID: {document_id}")
            return None, None, None, None, None, None, None

        try:
            # Initialize SmartDocumentProcessor
            doc_processor = SmartDocumentProcessor()
            
            # Process document and extract TOC
            chunks, content_map = doc_processor.process_document(markdown_content, resource_id=document_id)
            
            # Validate chunks
            if not chunks:
                logger.warning("No chunks generated, creating fallback chunk")
                chunks = [{
                    'document_id': document_id,
                    'document_title': document_info['title'],
                    'chapter_content': "Document content unavailable. Please check if:\n"
                                     "1. The document has readable text\n"
                                     "2. Contains more than 50 words\n"
                                     "3. Isn't image-based",
                    'chapter_title': 'Content Unavailable',
                    'chapter_no': 'N/A',
                    'chunk_index': 1,
                    'chunk_no': 1,
                    'total_chunks': 1,
                    'start_page_no': 1,
                    'end_page_no': 1
                }]
            
            logger.info(f"Created {len(chunks)} chunks from document")
            
            try:
                # Create embeddings
                embeddings = create_embeddings()
                
                # Prepare chunks for vector store
                texts = []
                metadatas = []
                
                for chunk in chunks:
                    # Ensure chunk has required content
                    if not chunk.get('chapter_content'):
                        continue
                        
                    # Add text and metadata
                    texts.append(chunk['chapter_content'])
                    metadatas.append({
                        'document_id': document_id,
                        'document_title': document_info['title'],
                        'chapter_no': chunk.get('chapter_no', 'N/A'),
                        'chapter_title': chunk.get('chapter_title', ''),
                        'chunk_index': chunk.get('chunk_index', 0),
                        'chunk_no': chunk.get('chunk_no', 0),
                        'total_chunks': chunk.get('total_chunks', 0),
                        'start_page_no': chunk.get('start_page_no'),
                        'end_page_no': chunk.get('end_page_no')
                    })
                
                if not texts:
                    logger.error("No valid chunks for vector store")
                    texts = ["Document content unavailable"]
                    metadatas = [{'document_id': document_id, 'fallback': True}]
                
                logger.info(f"Creating vector store with {len(texts)} chunks")
                
                # Create vector store
                vector_store = FAISS.from_texts(
                    texts,
                    embeddings,
                    metadatas=metadatas
                )
                
                # Create retriever
                retriever = LoggingVectorStoreRetriever(
                    vectorstore=vector_store,
                    search_kwargs={
                        "k": 4,
                        "fetch_k": 24,
                        "score_threshold": 0.5
                    }
                )
                
                # Save vector store
                vector_store_path = f"vector_stores/{document_id}"
                os.makedirs(os.path.dirname(vector_store_path), exist_ok=True)
                vector_store.save_local(vector_store_path)
                logger.info(f"Saved vector store for document ID: {document_id}")
                
                # Set up retrieval chain
                retrieval_chain = setup_llm_chain(retriever)
                
                # Store content map
                store_summary(document_id, content_map)
                
                # Generate summary
                try:
                    summary_prompt = f"""Analyze the structure and organization of this document based on its content map:

Document Title: {document_info['title']}

Content Map Structure:
{json.dumps(content_map, indent=2)}

Please provide:
1. Key points and main ideas evident from the document structure
2. Major themes and concepts based on the section organization
3. Document's organizational structure and hierarchy
4. Relevant keywords or tags derived from section titles

Focus your analysis on the organizational structure shown in the content map."""

                    initial_state = {
                        "input": summary_prompt,
                        "chat_history": [],
                        "document_title": document_info['title'],
                        "content_map": content_map
                    }
                    
                    conversation_app = create_conversation_graph(
                        retrieval_chain,
                        document_info['title'],
                        content_map
                    )
                    
                    config = {
                        "configurable": {
                            "thread_id": f"document_{document_id}",
                            "checkpoint_id": f"initial_summary_{document_id}",
                            "checkpoint_ns": "document_processing"
                        }
                    }
                    
                    summary_result = conversation_app.invoke(initial_state, config=config)
                    summary_and_tags = summary_result["answer"]
                    store_summary_and_tags(document_id, summary_and_tags)
                    mark_document_processed(document_id)
                    
                    return (vector_store, retrieval_chain, document_info, 
                            content_map, summary_and_tags, [], markdown_content)
                            
                except Exception as e:
                    logger.error(f"Error generating summary: {str(e)}")
                    return (vector_store, retrieval_chain, document_info, 
                            content_map, "", [], markdown_content)
                            
            except Exception as e:
                logger.error(f"Error creating vector store: {str(e)}")
                raise
                
        except Exception as e:
            logger.error(f"Error in document processing: {str(e)}")
            return None, None, None, None, None, None, None
            
    except Exception as e:
        logger.error(f"Unexpected error processing document ID {document_id}: {str(e)}")
        logger.error(traceback.format_exc())
        return None, None, None, None, None, None, None

def extract_text_from_document(file_path: str) -> Optional[str]:
    """Extract text content from various document types with enhanced Marker API support."""
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Handle plain text files directly
        if file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        # Get API key from environment
        api_key = os.getenv('MARKER_API_KEY')
        if not api_key:
            raise ValueError("Marker API key not found. Please set MARKER_API_KEY environment variable.")
        
        # Initialize Marker processor
        # marker_processor = MarkerDocumentProcessor(api_key)
        
        processor = MuPDFDocumentProcessor(output_dir="processed_documents")
        

        # Get document info for metadata
        document_info = get_document_info(file_path)
        
        # Process document with Marker
        markdown_content = processor.process_document(file_path, document_info)
        if not markdown_content:
            logger.error(f"Failed to process document: {file_path}")
            return None
            
        return markdown_content
        
    except Exception as e:
        logger.error(f"Error extracting text from document: {str(e)}")
        logger.error(traceback.format_exc())
        return None

def format_processed_text(text: str, document_info: Dict[str, Any], metrics: Dict[str, Any]) -> str:
    """Format processed text with metadata and metrics."""
    try:
        metadata_header = f"""
Title: {document_info['title']}
Type: {document_info.get('file_type', 'Unknown')}
Word Count: {metrics.get('word_count', 'N/A')}
Readability Score: {metrics.get('readability_score', 'N/A')}
Technical Terms: {', '.join(metrics.get('technical_terms', [])[:5])}
---
"""
        formatted_text = text.strip()
        
        # Add paragraph breaks if needed
        if '\n\n' not in formatted_text:
            formatted_text = re.sub(r'(?<!\n)\n(?!\n)', '\n\n', formatted_text)
            
        # Clean up whitespace
        formatted_text = re.sub(r'\s+', ' ', formatted_text)
        formatted_text = re.sub(r'\n\s*\n', '\n\n', formatted_text)
        
        return f"{metadata_header}\n{formatted_text}"
        
    except Exception as e:
        logger.error(f"Error formatting processed text: {str(e)}")
        return text

def process_query(retrieval_chain, query, document_id, document_title, content_map, session_identifier):
    """Process a user query about a document."""
    try:
        logger.info(f"Processing query for document {document_id} in session {session_identifier}")
        
        # Initialize chain logger
        chain_logger = ChainLogger()
        
        # Get chat history
        chat_history = ChatSession.objects.filter(
            session_identifier=session_identifier
        ).order_by('created_at')
        
        # Format the chat history for the LLM
        formatted_history = []
        for chat in chat_history:
            formatted_history.append(HumanMessage(content=chat.question))
            formatted_history.append(AIMessage(content=chat.answer))
        
        # Prepare chain input
        chain_input = {
            "input": query,
            "chat_history": formatted_history,
            "document_title": document_title,
            "content_map": content_map,
        }
        
        # Log the input for debugging
        logger.debug(f"Chain input for document {document_id}: {chain_input}")
        
        # Execute chain and get response
        response = retrieval_chain.invoke(chain_input)
        
        # Extract answer and context
        answer = response.get('answer', "I couldn't generate a response.")
        context = response.get('context', "")
        
        # Create retriever info without directly accessing the retriever
        retriever_info = {
            "chain_type": type(retrieval_chain).__name__,
            "query_timestamp": datetime.now().isoformat(),
            "document_id": document_id,
            "session_id": session_identifier,
            "context_length": len(context) if context else 0,
            "has_chat_history": bool(formatted_history)
        }
        
        # Log the chain execution
        log_file = chain_logger.log_chain_execution(
            query=query,
            chat_history=formatted_history,
            context=context,
            response=answer,
            document_title=document_title,
            content_map=content_map,
            retriever_info=retriever_info,
            query_id=f"{document_id}_{session_identifier}"
        )
        
        logger.info(f"Chain execution log saved to: {log_file}")
        
        return answer, formatted_history
        
    except Exception as e:
        error_msg = f"Error processing query: {str(e)}"
        logger.error(error_msg)
        logger.error(traceback.format_exc())
        return error_msg, []
    
def setup_llm_chain(retriever):
    """Set up the LLM chain with improved query handling."""
    
    # Enhanced system prompt for query contextualization
    contextualize_q_system_prompt = """Given the chat history and the latest user question about a document, your task is to:
1. Keep the original query intention clear and direct
2. Maintain any specific search terms or keywords from the original query
3. Add context from chat history only if it directly enhances the search
4. Do not make the query more generic or abstract
5. Do not ask clarifying questions - keep the reformulated query as a search statement

If the original query is clear and specific (like "show me table of contents" or "what does chapter 3 discuss"), 
use it as-is without reformulation.

Your output should be a search query, not a question to the user."""

    # Enhanced QA system prompt
    qa_system_prompt = """You are an AI assistant analyzing a document titled "{document_title}".
    
Use the provided content map, document context, and conversation history to give accurate answers.

Content Map:
{content_map}

Context from document:
{context}

Guidelines:
1. Use document content when available
2. Maintain conversation context without losing accuracy
3. If no relevant context is found, acknowledge this and suggest searching with different terms
4. For navigation requests (like table of contents), return full available information
5. Never speculate about unavailable content"""

    # Set up the contextualization prompt
    contextualize_q_prompt = ChatPromptTemplate.from_messages([
        ("system", contextualize_q_system_prompt),
        MessagesPlaceholder(variable_name="chat_history"),
        ("human", "{input}"),
    ])

    # Set up the QA prompt
    qa_prompt = ChatPromptTemplate.from_messages([
        ("system", qa_system_prompt),
        MessagesPlaceholder(variable_name="chat_history"),
        ("human", "{input}"),
    ])

    # Initialize LLM with lower temperature for more focused queries
    llm = ChatOpenAI(
        model="gpt-4o-mini",
        temperature=0.3,  # Lower temperature for more consistent query reformulation
    )

    # Create history-aware retriever
    history_aware_retriever = create_history_aware_retriever(
        llm,
        retriever,
        contextualize_q_prompt
    )

    # Create document chain
    doc_chain = create_stuff_documents_chain(
        llm,
        qa_prompt,
        document_variable_name="context"
    )

    # Create final retrieval chain
    return create_retrieval_chain(
        history_aware_retriever,
        doc_chain
    )